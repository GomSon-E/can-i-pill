"""Convert a Korean Markdown file to a formal DOCX report.

Supports:
- Headers (#, ##, ###, ####)
- Tables (| col | col |)
- Bullet lists (-) with complex nesting
- Numbered lists (1., 2., ...) with complex nesting
- Code blocks (``` ... ```)
- Bold (**text**) and italic (*text*)
- Inline code (`code`)
- Links ([text](url))
- Footnotes ([^1], [^1]: text)
- Horizontal rules (---)
- Plain paragraphs
"""
from __future__ import annotations

import argparse
import hashlib
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


# ========== Configuration Constants ==========
# MD 파일의 경로를 지정하는 상수 (필요에 따라 수정)
MD_INPUT_DIR = Path("../docs/v2")  # 마크다운 파일이 있는 디렉토리
MD_INPUT_FILE = "약-영양제-음식 상호작용 Agent 테스트 결과 보고서.md"  # 기본 파일명 (None이면 첫 번째 .md 파일 사용)

# 페이지 크기: A4 (한국 환경 기본) — 단위는 twip (1cm ≈ 567 twip)
PAGE_WIDTH_TWIP = 11906   # A4 width
PAGE_HEIGHT_TWIP = 16838  # A4 height
PAGE_MARGIN_LEFT_CM = 2.0
PAGE_MARGIN_RIGHT_CM = 2.0
PAGE_MARGIN_TOP_CM = 1.5
PAGE_MARGIN_BOTTOM_CM = 1.5
_TWIP_PER_CM = 567

# 본문에서 표가 차지할 수 있는 가용 폭 (twip)
USABLE_WIDTH_TWIP = (
    PAGE_WIDTH_TWIP
    - int(round(PAGE_MARGIN_LEFT_CM * _TWIP_PER_CM))
    - int(round(PAGE_MARGIN_RIGHT_CM * _TWIP_PER_CM))
)

HEADING_SIZES = {1: 20, 2: 16, 3: 13, 4: 12}
HEADING_COLORS = {1: "1F4E78", 2: "1F4E78", 3: "1F4E78", 4: "1F4E78"}  # 제목별 색상
BODY_FONT = "Malgun Gothic"
CODE_FONT = "Malgun Gothic"
BODY_SIZE = 11
TABLE_FONT_SIZE = 10
HEADER_FILL = "1F4E78"
LINK_COLOR = "0563C1"  # 링크 색상
LINK_UNDERLINE = True
CODE_BG = "F5F5F5"  # 코드 배경색
BLOCKQUOTE_BG = "EFF6FF"  # 인용문 배경색
BLOCKQUOTE_BORDER = "2563A8"  # 인용문 테두리 색상


INLINE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*\s][^*]*?)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+)`")
INLINE_LINK = re.compile(r"\[([^\[\]]+)\]\(([^)]+)\)")
INLINE_FOOTNOTE_REF = re.compile(r"\[\^(\w+)\]")
INLINE_FOOTNOTE_MARKER = re.compile(r"__FOOTNOTE_(\w+)__")
INLINE_HIGHLIGHT = re.compile(r"==(\S+(?:\s+\S+)*)?")  # Highlight ==text==
FOOTNOTE_DEF = re.compile(r"^\[\^(\w+)\]:\s+(.+)$", re.MULTILINE)
BLOCKQUOTE_PATTERN = re.compile(r"^>\s+(.+)$")


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


_HYPERLINK_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _github_slug(heading_text: str) -> str:
    """Compute a GitHub-style anchor slug from a heading's raw text.

    Mirrors the Markdown anchor that ``[text](#slug)`` would target so that
    a heading's bookmark name and an inline ``#slug`` link converge.
    """
    s = heading_text.strip().lower()
    # Strip inline markdown markers that don't survive into an anchor
    s = re.sub(r"\*+", "", s)
    s = re.sub(r"`+", "", s)
    # Drop punctuation except word chars, whitespace, and dashes (Unicode-aware)
    s = re.sub(r"[^\w\s-]", "", s, flags=re.UNICODE)
    s = re.sub(r"\s+", "-", s).strip("-")
    return s


def _anchor_to_bookmark_name(anchor_slug: str) -> str:
    """ASCII-only, ≤40 char bookmark name from a (possibly non-ASCII) slug.

    Internal-anchor links and heading bookmarks both pass through this so the
    names match regardless of whether the source had Korean characters.
    """
    canonical = anchor_slug.strip().lstrip("#").strip()
    digest = hashlib.sha1(canonical.encode("utf-8")).hexdigest()[:16]
    return f"bm_{digest}"


def _style_hyperlink_run(run: Any, base_font: str, base_size: float) -> None:
    run.font.name = base_font
    run.font.size = Pt(base_size)
    run.font.color.rgb = RGBColor(
        int(LINK_COLOR[0:2], 16),
        int(LINK_COLOR[2:4], 16),
        int(LINK_COLOR[4:6], 16),
    )
    if LINK_UNDERLINE:
        run.underline = True


def add_hyperlink(paragraph: Any, text: str, url: str, base_font: str = BODY_FONT,
                   base_size: float = BODY_SIZE) -> None:
    """Add a hyperlink run to a paragraph.

    - ``url`` starting with ``#`` → internal anchor: ``<w:hyperlink w:anchor="...">``.
      The bookmark name is derived from the same slug used by headings, so
      cross-references resolve when the heading actually exists in the doc.
    - Everything else → external link relationship.
    """
    run = paragraph.add_run(text)
    _style_hyperlink_run(run, base_font, base_size)

    hyperlink = OxmlElement("w:hyperlink")
    if url.startswith("#"):
        bookmark_name = _anchor_to_bookmark_name(url)
        hyperlink.set(qn("w:anchor"), bookmark_name)
    else:
        part = paragraph.part
        r_id = part.relate_to(url, _HYPERLINK_REL_TYPE, is_external=True)
        hyperlink.set(qn("r:id"), r_id)

    parent = run._element.getparent()
    parent.insert(parent.index(run._element), hyperlink)
    hyperlink.append(run._element)


def _add_heading_bookmark(heading_paragraph: Any, anchor_slug: str,
                           bookmark_id: int) -> None:
    """Wrap the heading paragraph content with bookmarkStart/End for the slug."""
    if not anchor_slug:
        return
    name = _anchor_to_bookmark_name(anchor_slug)
    p_element = heading_paragraph._element

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))

    # Insert bookmarkStart after pPr (if present), so it precedes the runs.
    pPr = p_element.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(bm_start)
    else:
        p_element.insert(0, bm_start)

    p_element.append(bm_end)


def _set_table_fixed_layout(table: Any, col_widths_twip: list[int]) -> None:
    """Force a fixed-width layout with explicit tblW, gridCol, and tcW values.

    한컴오피스 한워드는 ``<w:tblW>``/``<w:tcW>``/``<w:gridCol>`` 폭이 모두
    누락된 표를 만나면 1쪽 이후 렌더링에서 종료되는 사례가 확인되어, 모든
    표에 명시적 dxa(twip) 폭을 부여한다.
    """
    table.autofit = False
    table.allow_autofit = False
    # 참고: python-docx 1.x에는 ``table.layout_type`` 세터가 없으므로 아래에서
    # ``<w:tblLayout w:type="fixed"/>``를 직접 작성한다.

    total_twip = sum(col_widths_twip)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # tblW: total table width in dxa (twip)
    existing_tblW = tblPr.find(qn("w:tblW"))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_twip))
    tblW.set(qn("w:type"), "dxa")
    tblPr.insert(0, tblW)

    # tblLayout: fixed
    existing_tblLayout = tblPr.find(qn("w:tblLayout"))
    if existing_tblLayout is not None:
        tblPr.remove(existing_tblLayout)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # tblGrid: per-column widths
    existing_tblGrid = tbl.find(qn("w:tblGrid"))
    if existing_tblGrid is not None:
        tbl.remove(existing_tblGrid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_twip:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    # tblGrid follows tblPr
    tblPr.addnext(tblGrid)

    # Per-cell tcW
    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            existing_tcW = tcPr.find(qn("w:tcW"))
            if existing_tcW is not None:
                tcPr.remove(existing_tcW)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths_twip[col_idx]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _equal_column_widths(n_cols: int, total_twip: int = USABLE_WIDTH_TWIP) -> list[int]:
    base = total_twip // n_cols
    widths = [base] * n_cols
    widths[-1] += total_twip - base * n_cols  # absorb rounding remainder
    return widths


def add_blockquote(doc: Document, text: str, footnotes: dict[str, str] | None = None) -> None:
    """Add a blockquote paragraph with styling."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.4
    
    # 배경색
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), BLOCKQUOTE_BG)
    p._element.get_or_add_pPr().append(shading_elm)
    
    # 왼쪽 테두리 (컬러)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '24')
    left.set(qn('w:color'), BLOCKQUOTE_BORDER)
    pBdr.append(left)
    pPr.append(pBdr)
    
    apply_inline(p, text, base_italic=True, footnotes=footnotes)


def apply_inline(paragraph: Any, text: str, base_font: str = BODY_FONT,
                  base_size: float = BODY_SIZE, base_bold: bool = False,
                  base_italic: bool = False, footnotes: dict[str, str] | None = None) -> None:
    """Parse inline markdown (bold/italic/code/links/footnotes/highlight) and add as runs."""
    if footnotes is None:
        footnotes = {}
    
    # First, extract footnote references to replace later
    text_with_markers = text
    footnote_refs = {}
    for m in INLINE_FOOTNOTE_REF.finditer(text):
        ref_id = m.group(1)
        footnote_refs[f"__FOOTNOTE_{ref_id}__"] = ref_id
        text_with_markers = text_with_markers.replace(m.group(0), f"__FOOTNOTE_{ref_id}__", 1)
    
    tokens: list[tuple[str, Any]] = []
    i = 0
    text = text_with_markers
    while i < len(text):
        m_bold = INLINE_BOLD.search(text, i)
        m_italic = INLINE_ITALIC.search(text, i)
        m_code = INLINE_CODE.search(text, i)
        m_link = INLINE_LINK.search(text, i)
        m_footnote = INLINE_FOOTNOTE_MARKER.search(text, i)
        m_highlight = INLINE_HIGHLIGHT.search(text, i)

        candidates = []
        if m_bold:
            candidates.append((m_bold.start(), "bold", m_bold))
        if m_italic:
            candidates.append((m_italic.start(), "italic", m_italic))
        if m_code:
            candidates.append((m_code.start(), "code", m_code))
        if m_link:
            candidates.append((m_link.start(), "link", m_link))
        if m_footnote:
            candidates.append((m_footnote.start(), "footnote", m_footnote))
        if m_highlight:
            candidates.append((m_highlight.start(), "highlight", m_highlight))

        if not candidates or min(c[0] for c in candidates) != i:
            next_special = min((c[0] for c in candidates), default=len(text))
            plain = text[i:next_special]
            if plain:
                tokens.append(("plain", plain))
            i = next_special
        else:
            _, kind, m = next(c for c in candidates if c[0] == i)
            if kind == "link":
                link_text = m.group(1)
                link_url = m.group(2)
                tokens.append((kind, (link_text, link_url)))
            elif kind == "footnote":
                ref_id = m.group(1)
                tokens.append((kind, ref_id))
            elif kind == "highlight":
                highlighted_text = m.group(1)
                tokens.append((kind, highlighted_text))
            else:
                tokens.append((kind, m.group(1)))
            i = m.end()

    for kind, content in tokens:
        if kind == "code":
            run = paragraph.add_run(content)
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.italic = True
            run.font.color.rgb = RGBColor(220, 38, 38)
        elif kind == "link":
            link_text, link_url = content
            add_hyperlink(paragraph, link_text, link_url, base_font, base_size)
        elif kind == "footnote":
            ref_id = content
            run = paragraph.add_run(f"[{ref_id}]")
            run.font.name = base_font
            run.font.size = Pt(base_size - 2)
            run.superscript = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        elif kind == "highlight":
            run = paragraph.add_run(content)
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = RGBColor(220, 38, 38)
            run.bold = True
            # 배경 하이라이트
            shading_elm = OxmlElement('w:highlight')
            shading_elm.set(qn('w:val'), 'yellow')
            run._element.get_or_add_rPr().append(shading_elm)
        else:
            run = paragraph.add_run(content)
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.bold = base_bold or (kind == "bold")
            run.italic = base_italic or (kind == "italic")


def add_heading_paragraph(doc: Document, text: str, level: int,
                          footnotes: dict[str, str] | None = None,
                          bookmark_id: int | None = None) -> None:
    # 제목 앞에 한 줄 띄우기
    if level > 1:
        doc.add_paragraph()

    h = doc.add_heading("", level=min(level, 4))
    h.paragraph_format.space_after = Pt(10)  # 제목 아래 간격
    h.paragraph_format.space_before = Pt(0)

    apply_inline(h, text, base_size=HEADING_SIZES.get(level, 11), base_bold=True,
                 footnotes=footnotes)

    # 제목 색상 및 스타일
    color = HEADING_COLORS.get(level, "1F4E78")
    for run in h.runs:
        run.font.name = BODY_FONT
        run.bold = True
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        )

    # H2, H3에만 아래 보더라인 추가
    if level in [2, 3]:
        pPr = h._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # 내부 앵커 링크에서 참조 가능하도록 헤딩에 북마크 부착.
    if bookmark_id is not None:
        slug = _github_slug(text)
        if slug:
            _add_heading_bookmark(h, slug, bookmark_id)


def add_paragraph(doc: Document, text: str, footnotes: dict[str, str] | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5  # 줄 간격
    apply_inline(p, text, footnotes=footnotes)


def add_bullet(doc: Document, text: str, level: int = 0, 
               footnotes: dict[str, str] | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    apply_inline(p, text, footnotes=footnotes)


def add_numbered(doc: Document, text: str, level: int = 0,
                 footnotes: dict[str, str] | None = None) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    apply_inline(p, text, footnotes=footnotes)


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    
    # 배경색 추가
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), CODE_BG)
    p._element.get_or_add_pPr().append(shading_elm)
    
    # 테두리 추가
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'right', 'bottom']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)
    
    run = p.add_run(code.rstrip())
    run.font.name = CODE_FONT
    run.font.size = Pt(BODY_SIZE - 1)
    run.font.color.rgb = RGBColor(64, 64, 64)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")  # 더 굵은 선
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")  # 더 밝은 색
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_footnote_section(doc: Document, footnotes: dict[str, str]) -> None:
    """Add a footnote section at the end of the document."""
    if not footnotes:
        return
    
    # Add horizontal rule before footnotes
    add_horizontal_rule(doc)
    
    # Footnotes heading
    footnote_heading = doc.add_heading("각주", level=2)
    for run in footnote_heading.runs:
        run.font.name = BODY_FONT
        run.bold = True
    
    # Add each footnote
    for ref_id in sorted(footnotes.keys()):
        text = footnotes[ref_id]
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(6)
        
        # Add reference number as superscript at the beginning
        run_ref = p.add_run(f"{ref_id}")
        run_ref.font.size = Pt(BODY_SIZE - 2)
        run_ref.superscript = True
        run_ref.font.color.rgb = RGBColor(128, 128, 128)
        
        p.add_run(" ")
        apply_inline(p, text)


def parse_table_cells(line: str) -> list[str]:
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def is_separator_row(line: str) -> bool:
    cells = parse_table_cells(line)
    return all(re.match(r":?-{2,}:?$", c) for c in cells if c)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              footnotes: dict[str, str] | None = None) -> None:
    if not headers:
        return

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = _equal_column_widths(n_cols)

    # 헤더 행
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_inline(para, header, base_size=TABLE_FONT_SIZE, base_bold=True,
                     footnotes=footnotes)
        for run in para.runs:
            run.font.name = BODY_FONT
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "2563A8")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 데이터 행
    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        for col_idx in range(n_cols):
            value = row[col_idx] if col_idx < len(row) else ""
            cells[col_idx].text = ""
            para = cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_inline(para, value, base_size=TABLE_FONT_SIZE, footnotes=footnotes)
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 모든 셀이 채워진 후에 고정 폭(dxa) 부여 — autofit/0-width 표가
    # 한컴오피스 한워드에서 렌더링 중 종료되는 문제 회피.
    _set_table_fixed_layout(table, col_widths)

    doc.add_paragraph()


def get_indent_level(line: str) -> int:
    """Calculate indent level of a line (2 spaces = 1 level)."""
    indent = len(line) - len(line.lstrip())
    return indent // 2


def parse_list_item(line: str) -> tuple[str, int, str] | None:
    """Parse list item. Returns (text, level, type) or None."""
    stripped = line.rstrip("\n")
    
    # Bullet list
    m_bullet = re.match(r"^(\s*)-\s+(.+)$", stripped)
    if m_bullet:
        indent = len(m_bullet.group(1))
        level = indent // 2
        text = m_bullet.group(2)
        return (text, level, "bullet")
    
    # Numbered list
    m_number = re.match(r"^(\s*)\d+\.\s+(.+)$", stripped)
    if m_number:
        indent = len(m_number.group(1))
        level = indent // 2
        text = m_number.group(2)
        return (text, level, "number")
    
    return None


def render_lines(doc: Document, lines: list[str], footnotes: dict[str, str] | None = None) -> None:
    if footnotes is None:
        footnotes = {}

    bookmark_counter = 0

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip("\n")

        # Code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].rstrip("\n").startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, "\n".join(code_lines))
            continue

        # Empty line
        if not stripped.strip():
            i += 1
            continue

        # Blockquote
        m_blockquote = BLOCKQUOTE_PATTERN.match(stripped)
        if m_blockquote:
            blockquote_lines = [m_blockquote.group(1)]
            i += 1
            while i < n:
                nxt = lines[i].rstrip("\n")
                m_next = BLOCKQUOTE_PATTERN.match(nxt)
                if not m_next:
                    break
                blockquote_lines.append(m_next.group(1))
                i += 1
            add_blockquote(doc, " ".join(blockquote_lines), footnotes)
            continue
        
        # Horizontal rule
        if re.match(r"^\s*-{3,}\s*$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            bookmark_counter += 1
            add_heading_paragraph(doc, text, level, footnotes,
                                   bookmark_id=bookmark_counter)
            i += 1
            continue

        # Table
        if stripped.lstrip().startswith("|"):
            header_line = stripped
            if i + 1 < n and lines[i + 1].lstrip().startswith("|") and is_separator_row(lines[i + 1]):
                headers = parse_table_cells(header_line)
                i += 2
                rows = []
                while i < n and lines[i].lstrip().startswith("|"):
                    cells = parse_table_cells(lines[i].rstrip("\n"))
                    rows.append(cells)
                    i += 1
                add_table(doc, headers, rows, footnotes)
                continue

        # Complex nested lists
        list_item = parse_list_item(line)
        if list_item:
            item_texts, item_level, item_type = list_item
            list_items = [(item_texts, item_level, item_type)]
            i += 1
            
            # Collect consecutive list items
            while i < n:
                nxt = lines[i].rstrip("\n")
                next_is_list = parse_list_item(nxt)
                if not next_is_list:
                    break
                next_text, next_level, next_type = next_is_list
                list_items.append((next_text, next_level, next_type))
                i += 1
            
            # Render accumulated list items
            for text, level, ltype in list_items:
                if ltype == "bullet":
                    add_bullet(doc, text, level=level, footnotes=footnotes)
                else:
                    add_numbered(doc, text, level=level, footnotes=footnotes)
            continue

        # Plain paragraph (collect consecutive non-empty/non-special lines)
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip("\n")
            if not nxt.strip():
                break
            if nxt.startswith("#") or nxt.startswith("```") or nxt.lstrip().startswith("|"):
                break
            if re.match(r"^\s*-{3,}\s*$", nxt):
                break
            if parse_list_item(nxt):
                break
            para_lines.append(nxt)
            i += 1
        add_paragraph(doc, " ".join(para_lines), footnotes=footnotes)


def build_document(md_text: str, title: str) -> Document:
    doc = Document()

    section = doc.sections[0]
    # A4 (한국 환경 기본). 표 폭 계산도 이 값에 맞춰 있음.
    section.page_width = Twips(PAGE_WIDTH_TWIP)
    section.page_height = Twips(PAGE_HEIGHT_TWIP)
    section.left_margin = Cm(PAGE_MARGIN_LEFT_CM)
    section.right_margin = Cm(PAGE_MARGIN_RIGHT_CM)
    section.top_margin = Cm(PAGE_MARGIN_TOP_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_BOTTOM_CM)

    # Extract footnotes from the markdown
    footnotes: dict[str, str] = {}
    for m in FOOTNOTE_DEF.finditer(md_text):
        ref_id = m.group(1)
        footnote_text = m.group(2).strip()
        footnotes[ref_id] = footnote_text
    
    # Remove footnote definitions from the text
    md_text_clean = FOOTNOTE_DEF.sub("", md_text)

    # 표지
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = BODY_FONT
    doc.add_paragraph()

    # 첫 H1을 표지로 흡수하므로 본문 렌더에서는 첫 H1 스킵
    lines = md_text_clean.splitlines(keepends=True)
    # Strip BOM if present
    if lines and lines[0].startswith("﻿"):
        lines[0] = lines[0].lstrip("﻿")

    # Skip the leading H1 (title) if it matches title or is the first non-empty line
    j = 0
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j < len(lines) and lines[j].lstrip().startswith("# "):
        j += 1
    lines = lines[j:]

    render_lines(doc, lines, footnotes)
    
    # Add footnote section at the end
    if footnotes:
        add_footnote_section(doc, footnotes)
    
    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert a Markdown file to formal DOCX.")
    parser.add_argument("input", nargs="?", help="Input .md file (defaults to MD_INPUT_FILE if not specified)")
    parser.add_argument("output", nargs="?", help="Output .docx file")
    parser.add_argument("--title", default=None, help="Document title (defaults to first H1)")
    parser.add_argument("--input-dir", type=Path, default=MD_INPUT_DIR, 
                        help=f"Directory containing markdown files (default: {MD_INPUT_DIR})")
    args = parser.parse_args()

    # Determine input and output paths
    if args.input is None:
        # If no input provided, use MD_INPUT_FILE or first .md file
        input_dir = args.input_dir.resolve()
        if not input_dir.exists():
            print(f"Input directory not found: {input_dir}", file=sys.stderr)
            print(f"Configure MD_INPUT_DIR or provide an input file.", file=sys.stderr)
            return 1
        
        # Use MD_INPUT_FILE if specified, otherwise use first .md file
        if MD_INPUT_FILE:
            input_path = input_dir / MD_INPUT_FILE
        else:
            md_files = list(input_dir.glob("*.md"))
            if not md_files:
                print(f"No .md files found in {input_dir}", file=sys.stderr)
                return 1
            input_path = md_files[0]
        
        print(f"Using default file: {input_path}")
    else:
        input_path = Path(args.input)
    
    if args.output is None:
        output_path = input_path.with_suffix(".docx")
    else:
        output_path = Path(args.output)

    if not input_path.exists():
        print(f"Input not found: {input_path}", file=sys.stderr)
        return 1

    md_text = input_path.read_text(encoding="utf-8")

    if args.title is None:
        m = re.search(r"^#\s+(.+)$", md_text, re.MULTILINE)
        title = m.group(1).strip() if m else input_path.stem
    else:
        title = args.title

    doc = build_document(md_text, title)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
